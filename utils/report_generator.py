"""
Modulo para generacion de informes completos de auditoria RCF
"""

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib import colors
from datetime import datetime
import io
from config.settings import CONFIGURACION_INFORME, CONFIGURACION


def set_cell_shading(cell, color):
    """Aplica color de fondo a una celda de Word"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_to_doc(doc, df, title=None, max_rows=20):
    """Agrega una tabla formateada al documento Word"""
    if df is None or len(df) == 0:
        return

    if title:
        doc.add_heading(title, level=3)

    # Limitar filas si es necesario
    df_display = df.head(max_rows) if len(df) > max_rows else df

    # Crear tabla
    table = doc.add_table(rows=1, cols=len(df_display.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados
    header_cells = table.rows[0].cells
    for i, col in enumerate(df_display.columns):
        header_cells[i].text = str(col)
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '0066CC')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Datos
    for _, row in df_display.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            if pd.isna(value):
                row_cells[i].text = ''
            elif isinstance(value, float):
                row_cells[i].text = f'{value:,.2f}'
            elif isinstance(value, int):
                row_cells[i].text = f'{value:,}'
            elif isinstance(value, pd.Timestamp):
                row_cells[i].text = value.strftime('%d/%m/%Y')
            else:
                row_cells[i].text = str(value)[:50]  # Truncar texto largo

    if len(df) > max_rows:
        doc.add_paragraph(f'(Mostrando {max_rows} de {len(df)} registros)')

    doc.add_paragraph()  # Espacio


def generar_informe_word(datos: dict, analisis: dict) -> bytes:
    """
    Genera un informe completo en formato Word con todas las tablas y analisis
    """
    doc = Document()

    # Configurar estilos
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ============================================
    # PORTADA
    # ============================================
    doc.add_paragraph()
    doc.add_paragraph()

    titulo = doc.add_heading(CONFIGURACION_INFORME['titulo'], 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_heading(CONFIGURACION_INFORME['subtitulo'], 1)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Informacion general centrada
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Entidad: ').bold = True
    info.add_run(CONFIGURACION['nombre_entidad'] + '\n\n')
    info.add_run('Ejercicio auditado: ').bold = True
    info.add_run(CONFIGURACION['ejercicio_auditado'] + '\n\n')
    info.add_run('Fecha del informe: ').bold = True
    info.add_run(datetime.now().strftime('%d/%m/%Y'))

    doc.add_page_break()

    # ============================================
    # INDICE
    # ============================================
    doc.add_heading('INDICE', 1)

    indice_items = [
        '1. Resumen Ejecutivo',
        '2. Introduccion y Marco Normativo',
        '3. Analisis de Facturas en Papel (Seccion V.1)',
        '4. Tiempos de Anotacion en RCF (Seccion V.2)',
        '5. Validaciones de Contenido HAP/1650/2015 (Seccion V.3)',
        '6. Tramitacion de Facturas (Seccion V.4)',
        '7. Obligaciones de Control y Morosidad (Seccion V.5)',
        '8. Conclusiones y Recomendaciones'
    ]

    for item in indice_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # ============================================
    # 1. RESUMEN EJECUTIVO
    # ============================================
    doc.add_heading('1. RESUMEN EJECUTIVO', 1)

    doc.add_paragraph(f"""
El presente informe recoge los resultados de la auditoria del Registro Contable de Facturas (RCF)
de {CONFIGURACION['nombre_entidad']} correspondiente al ejercicio {CONFIGURACION['ejercicio_auditado']},
en cumplimiento del articulo 12.3 de la Ley 25/2013, de 27 de diciembre, de impulso de la factura
electronica y creacion del registro contable de facturas en el Sector Publico.
    """)

    # Metricas principales
    doc.add_heading('1.1. Cifras Principales', 2)

    total_facturas = len(datos['rcf'])
    facturas_electronicas = len(datos['rcf'][datos['rcf']['es_papel'] == False])
    facturas_papel = len(datos['rcf'][datos['rcf']['es_papel'] == True])

    tabla_resumen = [
        ['Concepto', 'Valor'],
        ['Total de facturas en RCF', f'{total_facturas:,}'],
        ['Facturas electronicas', f'{facturas_electronicas:,}'],
        ['Facturas en papel', f'{facturas_papel:,}'],
        ['Porcentaje electronicas', f'{(facturas_electronicas/total_facturas*100):.1f}%' if total_facturas > 0 else '0%'],
        ['Solicitudes de anulacion', f'{len(datos["anulaciones"]):,}'],
    ]

    table = doc.add_table(rows=len(tabla_resumen), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, fila in enumerate(tabla_resumen):
        for j, valor in enumerate(fila):
            table.rows[i].cells[j].text = str(valor)
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Hallazgos principales
    doc.add_heading('1.2. Principales Hallazgos', 2)

    hallazgos = []

    if 'facturas_papel' in analisis:
        papel = analisis['facturas_papel']
        if papel.get('total_sospechosas', 0) > 0:
            hallazgos.append(f"Se han identificado {papel['total_sospechosas']:,} facturas en papel susceptibles de incumplir la obligatoriedad de factura electronica.")

    if 'anotacion' in analisis:
        anot = analisis['anotacion']
        if anot.get('facturas_retenidas', 0) > 0:
            hallazgos.append(f"ALERTA: Se han detectado {anot['facturas_retenidas']} facturas retenidas en FACe pendientes de descarga.")
        hallazgos.append(f"El tiempo medio de anotacion en RCF es de {anot.get('tiempo_medio_min', 0):.2f} minutos.")

    if 'validaciones' in analisis:
        val = analisis['validaciones']
        hallazgos.append(f"El porcentaje de cumplimiento de validaciones HAP/1650/2015 es del {val.get('porcentaje_cumplimiento', 100):.2f}%.")

    if 'obligaciones' in analisis:
        oblig = analisis['obligaciones']
        if oblig.get('facturas_3_meses', 0) > 0:
            hallazgos.append(f"Existen {oblig['facturas_3_meses']} facturas con mas de 3 meses sin reconocimiento de obligacion.")

    for hallazgo in hallazgos:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(hallazgo)

    doc.add_page_break()

    # ============================================
    # 2. INTRODUCCION Y MARCO NORMATIVO
    # ============================================
    doc.add_heading('2. INTRODUCCION Y MARCO NORMATIVO', 1)

    doc.add_heading('2.1. Marco Legal Aplicable', 2)

    doc.add_paragraph("""
La presente auditoria se realiza en cumplimiento de la siguiente normativa:

- Ley 25/2013, de 27 de diciembre, de impulso de la factura electronica y creacion del registro contable de facturas en el Sector Publico.
- Orden HAP/492/2014, de 27 de marzo, por la que se regulan los requisitos funcionales y tecnicos del registro contable de facturas.
- Orden HAP/1650/2015, de 31 de julio, por la que se modifican la Orden HAP/492/2014 y la Orden HAP/1074/2014.
- Ley 15/2010, de 5 de julio, de modificacion de la Ley 3/2004 por la que se establecen medidas de lucha contra la morosidad.
    """)

    doc.add_heading('2.2. Objetivos de la Auditoria', 2)

    objetivos = [
        'Verificar el cumplimiento de la obligatoriedad de factura electronica.',
        'Analizar los tiempos de anotacion en el RCF.',
        'Comprobar la aplicacion de las validaciones de la Orden HAP/1650/2015.',
        'Revisar la tramitacion de solicitudes de anulacion.',
        'Controlar las facturas pendientes de reconocimiento de obligacion.',
    ]

    for obj in objetivos:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(obj)

    doc.add_heading('2.3. Archivos Analizados', 2)

    archivos_tabla = [
        ['Archivo', 'Registros'],
        ['Registro Contable de Facturas (RCF)', f'{len(datos["rcf"]):,}'],
        ['Facturas FACe', f'{len(datos["face"]):,}'],
        ['Solicitudes de Anulacion', f'{len(datos["anulaciones"]):,}'],
        ['Historico de Estados', f'{len(datos["estados"]):,}'],
    ]

    table = doc.add_table(rows=len(archivos_tabla), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, fila in enumerate(archivos_tabla):
        for j, valor in enumerate(fila):
            table.rows[i].cells[j].text = str(valor)
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ============================================
    # 3. FACTURAS EN PAPEL
    # ============================================
    doc.add_heading('3. ANALISIS DE FACTURAS EN PAPEL (Seccion V.1)', 1)

    doc.add_paragraph("""
Segun el articulo 4 de la Ley 25/2013, la factura electronica es obligatoria para las personas juridicas
y para facturas de importe superior a 3.000 euros (base imponible). Se analizan las facturas en papel
que podrian estar incumpliendo esta obligatoriedad.
    """)

    if 'facturas_papel' in analisis:
        papel = analisis['facturas_papel']

        doc.add_heading('3.1. Resumen de Resultados', 2)

        resumen_papel = [
            ['Concepto', 'Valor'],
            ['Total facturas en papel', f'{papel.get("total_papel", 0):,}'],
            ['Susceptibles de incumplimiento', f'{papel.get("total_sospechosas", 0):,}'],
            ['Importe total sospechoso', f'{papel.get("importe_sospechoso", 0):,.2f} EUR'],
        ]

        table = doc.add_table(rows=len(resumen_papel), cols=2)
        table.style = 'Light Grid Accent 1'
        for i, fila in enumerate(resumen_papel):
            for j, valor in enumerate(fila):
                table.rows[i].cells[j].text = str(valor)
                if i == 0:
                    table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Top 10 por importe
        if 'top_10_importe' in papel and len(papel['top_10_importe']) > 0:
            add_table_to_doc(doc, papel['top_10_importe'], '3.2. Top 10 Facturas por Importe', 10)

        # Top proveedores
        if 'top_proveedores' in papel and len(papel['top_proveedores']) > 0:
            add_table_to_doc(doc, papel['top_proveedores'], '3.3. Top 10 Proveedores por Base Imponible Acumulada', 10)

        # Ranking OC
        if 'ranking_oc' in papel and len(papel['ranking_oc']) > 0:
            add_table_to_doc(doc, papel['ranking_oc'], '3.4. Ranking de Oficinas Contables', 10)

        # Ranking UT
        if 'ranking_ut' in papel and len(papel['ranking_ut']) > 0:
            add_table_to_doc(doc, papel['ranking_ut'], '3.5. Ranking de Unidades Tramitadoras', 10)

        # Evolucion mensual
        if 'evolucion_mensual' in papel and len(papel['evolucion_mensual']) > 0:
            add_table_to_doc(doc, papel['evolucion_mensual'].reset_index(), '3.6. Evolucion Mensual', 12)
    else:
        doc.add_paragraph('No se ha realizado el analisis de facturas en papel. Navegue por la seccion correspondiente de la aplicacion.')

    doc.add_page_break()

    # ============================================
    # 4. TIEMPOS DE ANOTACION
    # ============================================
    doc.add_heading('4. TIEMPOS DE ANOTACION EN RCF (Seccion V.2)', 1)

    doc.add_paragraph("""
Se analiza el tiempo transcurrido desde que una factura se registra en FACe hasta su anotacion en el RCF.
Tiempos excesivos pueden indicar problemas en la integracion o retencion indebida de facturas.
    """)

    if 'anotacion' in analisis:
        anot = analisis['anotacion']

        doc.add_heading('4.1. Facturas Retenidas en FACe', 2)

        if anot.get('facturas_retenidas', 0) > 0:
            p = doc.add_paragraph()
            p.add_run('ALERTA: ').bold = True
            p.add_run(f'Se han detectado {anot["facturas_retenidas"]} facturas retenidas en FACe que no han sido descargadas al RCF.')

            if 'df_retenidas' in anot and len(anot['df_retenidas']) > 0:
                add_table_to_doc(doc, anot['df_retenidas'], 'Detalle de Facturas Retenidas', 20)
        else:
            doc.add_paragraph('No se han detectado facturas retenidas. Todas las facturas de FACe estan correctamente anotadas en el RCF.')

        doc.add_heading('4.2. Estadisticas de Tiempos de Anotacion', 2)

        tiempos_tabla = [
            ['Metrica', 'Valor'],
            ['Tiempo medio', f'{anot.get("tiempo_medio_min", 0):.2f} minutos'],
            ['Tiempo mediano', f'{anot.get("tiempo_mediano_min", 0):.2f} minutos'],
            ['Tiempo minimo', f'{anot.get("tiempo_min_min", 0):.2f} minutos'],
            ['Tiempo maximo', f'{anot.get("tiempo_max_min", 0):.2f} minutos'],
        ]

        table = doc.add_table(rows=len(tiempos_tabla), cols=2)
        table.style = 'Light Grid Accent 1'
        for i, fila in enumerate(tiempos_tabla):
            for j, valor in enumerate(fila):
                table.rows[i].cells[j].text = str(valor)
                if i == 0:
                    table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Top demoras
        if 'top_demoras' in anot and len(anot['top_demoras']) > 0:
            add_table_to_doc(doc, anot['top_demoras'], '4.3. Top 10 Facturas con Mayor Demora', 10)

        # Stats mensuales
        if 'stats_mensuales' in anot and len(anot['stats_mensuales']) > 0:
            add_table_to_doc(doc, anot['stats_mensuales'], '4.4. Estadisticas Mensuales de Tiempos', 12)

        # Ranking OC tiempos
        if 'ranking_oc_tiempos' in anot and len(anot['ranking_oc_tiempos']) > 0:
            add_table_to_doc(doc, anot['ranking_oc_tiempos'], '4.5. Ranking de Oficinas Contables por Tiempo Medio', 10)
    else:
        doc.add_paragraph('No se ha realizado el analisis de tiempos de anotacion. Navegue por la seccion correspondiente de la aplicacion.')

    doc.add_page_break()

    # ============================================
    # 5. VALIDACIONES HAP/1650/2015
    # ============================================
    doc.add_heading('5. VALIDACIONES DE CONTENIDO HAP/1650/2015 (Seccion V.3)', 1)

    doc.add_paragraph("""
La Orden HAP/1650/2015 establece 8 validaciones obligatorias para las facturas electronicas con fecha
de expedicion posterior al 15 de octubre de 2015. Se analiza el cumplimiento de cada una de ellas.
    """)

    if 'validaciones' in analisis:
        val = analisis['validaciones']

        doc.add_heading('5.1. Resumen de Cumplimiento', 2)

        resumen_val = [
            ['Concepto', 'Valor'],
            ['Facturas validadas', f'{val.get("total_facturas_validadas", 0):,}'],
            ['Total incumplimientos', f'{val.get("total_incumplimientos", 0):,}'],
            ['Porcentaje cumplimiento', f'{val.get("porcentaje_cumplimiento", 100):.2f}%'],
        ]

        table = doc.add_table(rows=len(resumen_val), cols=2)
        table.style = 'Light Grid Accent 1'
        for i, fila in enumerate(resumen_val):
            for j, valor in enumerate(fila):
                table.rows[i].cells[j].text = str(valor)
                if i == 0:
                    table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        doc.add_heading('5.2. Detalle por Validacion', 2)

        if 'validaciones' in val:
            tabla_val = [['Codigo', 'Validacion', 'Incumplimientos', 'Porcentaje']]
            for codigo, resultado in val['validaciones'].items():
                tabla_val.append([
                    codigo,
                    resultado['nombre'][:40],
                    str(resultado['num_incumplimientos']),
                    f"{resultado['porcentaje']:.2f}%"
                ])

            table = doc.add_table(rows=len(tabla_val), cols=4)
            table.style = 'Light Grid Accent 1'
            for i, fila in enumerate(tabla_val):
                for j, valor in enumerate(fila):
                    table.rows[i].cells[j].text = str(valor)
                    if i == 0:
                        table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Analisis de rechazos
        if 'rechazos' in analisis:
            rechazos = analisis['rechazos']
            doc.add_heading('5.3. Analisis de Facturas Rechazadas', 2)

            doc.add_paragraph(f"Total facturas rechazadas: {rechazos.get('total_rechazadas', 0):,} ({rechazos.get('porcentaje', 0):.2f}%)")

            if rechazos.get('por_motivo'):
                doc.add_heading('Principales Motivos de Rechazo', 3)
                motivos_list = sorted(rechazos['por_motivo'].items(), key=lambda x: x[1], reverse=True)[:10]

                tabla_motivos = [['Motivo', 'Cantidad']]
                for motivo, cantidad in motivos_list:
                    tabla_motivos.append([str(motivo)[:60], str(cantidad)])

                table = doc.add_table(rows=len(tabla_motivos), cols=2)
                table.style = 'Light Grid Accent 1'
                for i, fila in enumerate(tabla_motivos):
                    for j, valor in enumerate(fila):
                        table.rows[i].cells[j].text = str(valor)
                        if i == 0:
                            table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    else:
        doc.add_paragraph('No se ha realizado el analisis de validaciones. Navegue por la seccion correspondiente de la aplicacion.')

    doc.add_page_break()

    # ============================================
    # 6. TRAMITACION DE FACTURAS
    # ============================================
    doc.add_heading('6. TRAMITACION DE FACTURAS (Seccion V.4)', 1)

    doc.add_paragraph("""
Se analiza la tramitacion de las facturas, incluyendo las solicitudes de anulacion, la evolucion
de estados y el reconocimiento de obligaciones.
    """)

    if 'tramitacion' in analisis:
        tram = analisis['tramitacion']

        doc.add_heading('6.1. Solicitudes de Anulacion', 2)

        anulaciones_tabla = [
            ['Concepto', 'Valor'],
            ['Total solicitudes', f'{tram.get("total_anulaciones", 0):,}'],
            ['Anulaciones aceptadas', f'{tram.get("anulaciones_aceptadas", 0):,}'],
            ['Con comentario', f'{tram.get("anulaciones_con_comentario", 0):,}'],
        ]

        table = doc.add_table(rows=len(anulaciones_tabla), cols=2)
        table.style = 'Light Grid Accent 1'
        for i, fila in enumerate(anulaciones_tabla):
            for j, valor in enumerate(fila):
                table.rows[i].cells[j].text = str(valor)
                if i == 0:
                    table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Detalle anulaciones
        if 'detalle_anulaciones' in tram and len(tram['detalle_anulaciones']) > 0:
            add_table_to_doc(doc, tram['detalle_anulaciones'], '6.2. Detalle de Solicitudes de Anulacion', 20)

        doc.add_heading('6.3. Distribucion de Estados', 2)

        # Distribucion estados
        if 'distribucion_estados' in tram and len(tram['distribucion_estados']) > 0:
            add_table_to_doc(doc, tram['distribucion_estados'], None, 15)

        doc.add_heading('6.4. Reconocimiento de Obligacion y Pagos', 2)

        pagos_tabla = [
            ['Concepto', 'Valor'],
            ['Facturas pagadas', f'{tram.get("facturas_pagadas", 0):,}'],
            ['Importe total pagado', f'{tram.get("importe_pagado", 0):,.2f} EUR'],
            ['Facturas contabilizadas', f'{tram.get("facturas_contabilizadas", 0):,}'],
        ]

        table = doc.add_table(rows=len(pagos_tabla), cols=2)
        table.style = 'Light Grid Accent 1'
        for i, fila in enumerate(pagos_tabla):
            for j, valor in enumerate(fila):
                table.rows[i].cells[j].text = str(valor)
                if i == 0:
                    table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

        # Secuencias de estados
        if 'secuencias_estados' in tram and len(tram['secuencias_estados']) > 0:
            add_table_to_doc(doc, tram['secuencias_estados'], '6.5. Secuencias de Estados mas Frecuentes', 10)
    else:
        doc.add_paragraph('No se ha realizado el analisis de tramitacion. Navegue por la seccion correspondiente de la aplicacion.')

    doc.add_page_break()

    # ============================================
    # 7. OBLIGACIONES Y MOROSIDAD
    # ============================================
    doc.add_heading('7. OBLIGACIONES DE CONTROL Y MOROSIDAD (Seccion V.5)', 1)

    doc.add_paragraph("""
Segun el articulo 12.3 de la Ley 25/2013, se deben controlar trimestralmente las facturas pendientes
de reconocimiento de obligacion. Se analizan especialmente las facturas con mas de 3 meses de antiguedad.
    """)

    if 'obligaciones' in analisis:
        oblig = analisis['obligaciones']

        doc.add_heading('7.1. Facturas Pendientes mas de 3 Meses', 2)

        if oblig.get('facturas_3_meses', 0) > 0:
            p = doc.add_paragraph()
            p.add_run('ALERTA: ').bold = True
            p.add_run(f'Se han identificado {oblig["facturas_3_meses"]} facturas con mas de 3 meses sin reconocimiento de obligacion.')

            pendientes_tabla = [
                ['Concepto', 'Valor'],
                ['Facturas pendientes >3 meses', f'{oblig.get("facturas_3_meses", 0):,}'],
                ['Importe total pendiente', f'{oblig.get("importe_pendiente", 0):,.2f} EUR'],
                ['Dias medio pendiente', f'{oblig.get("dias_medio_pendiente", 0):.0f} dias'],
                ['Antiguedad maxima', f'{oblig.get("dias_max_pendiente", 0):.0f} dias'],
            ]

            table = doc.add_table(rows=len(pendientes_tabla), cols=2)
            table.style = 'Light Grid Accent 1'
            for i, fila in enumerate(pendientes_tabla):
                for j, valor in enumerate(fila):
                    table.rows[i].cells[j].text = str(valor)
                    if i == 0:
                        table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

            doc.add_paragraph()

            # Detalle pendientes
            if 'detalle_pendientes' in oblig and len(oblig['detalle_pendientes']) > 0:
                add_table_to_doc(doc, oblig['detalle_pendientes'], '7.2. Detalle de Facturas Pendientes', 30)

            # Ranking OC pendientes
            if 'ranking_oc_pendientes' in oblig and len(oblig['ranking_oc_pendientes']) > 0:
                add_table_to_doc(doc, oblig['ranking_oc_pendientes'], '7.3. Ranking de Oficinas Contables con Pendientes', 10)

            # Ranking UT pendientes
            if 'ranking_ut_pendientes' in oblig and len(oblig['ranking_ut_pendientes']) > 0:
                add_table_to_doc(doc, oblig['ranking_ut_pendientes'], '7.4. Ranking de Unidades Tramitadoras con Pendientes', 10)
        else:
            doc.add_paragraph('No se han detectado facturas con mas de 3 meses pendientes de reconocimiento de obligacion.')

        # Morosidad
        if oblig.get('morosidad') and oblig['morosidad'].get('total_con_retraso', 0) > 0:
            doc.add_heading('7.5. Analisis de Morosidad', 2)

            morosidad = oblig['morosidad']
            morosidad_tabla = [
                ['Concepto', 'Valor'],
                ['Facturas pagadas con retraso (>30 dias)', f'{morosidad.get("total_con_retraso", 0):,}'],
                ['Dias medio de retraso', f'{morosidad.get("dias_retraso_medio", 0):.0f} dias'],
                ['Retraso maximo', f'{morosidad.get("dias_retraso_max", 0):.0f} dias'],
            ]

            table = doc.add_table(rows=len(morosidad_tabla), cols=2)
            table.style = 'Light Grid Accent 1'
            for i, fila in enumerate(morosidad_tabla):
                for j, valor in enumerate(fila):
                    table.rows[i].cells[j].text = str(valor)
                    if i == 0:
                        table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    else:
        doc.add_paragraph('No se ha realizado el analisis de obligaciones. Navegue por la seccion correspondiente de la aplicacion.')

    doc.add_page_break()

    # ============================================
    # 8. CONCLUSIONES Y RECOMENDACIONES
    # ============================================
    doc.add_heading('8. CONCLUSIONES Y RECOMENDACIONES', 1)

    doc.add_heading('8.1. Conclusiones', 2)

    conclusiones = []

    conclusiones.append('El Registro Contable de Facturas cumple en general con los requisitos establecidos en la Ley 25/2013 y su normativa de desarrollo.')

    if 'facturas_papel' in analisis and analisis['facturas_papel'].get('total_sospechosas', 0) > 0:
        conclusiones.append(f"Se han identificado {analisis['facturas_papel']['total_sospechosas']:,} facturas en papel que podrian incumplir la obligatoriedad de factura electronica, por un importe total de {analisis['facturas_papel'].get('importe_sospechoso', 0):,.2f} EUR.")

    if 'anotacion' in analisis and analisis['anotacion'].get('facturas_retenidas', 0) > 0:
        conclusiones.append(f"Existen {analisis['anotacion']['facturas_retenidas']} facturas retenidas en FACe que requieren investigacion inmediata.")

    if 'validaciones' in analisis:
        porc = analisis['validaciones'].get('porcentaje_cumplimiento', 100)
        if porc < 100:
            conclusiones.append(f"El porcentaje de cumplimiento de las validaciones HAP/1650/2015 es del {porc:.2f}%, siendo necesario revisar los incumplimientos detectados.")

    if 'obligaciones' in analisis and analisis['obligaciones'].get('facturas_3_meses', 0) > 0:
        conclusiones.append(f"Se han detectado {analisis['obligaciones']['facturas_3_meses']} facturas con mas de 3 meses pendientes de reconocimiento de obligacion, lo que podria suponer un incumplimiento de los plazos legales.")

    for i, conclusion in enumerate(conclusiones, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(conclusion)

    doc.add_heading('8.2. Recomendaciones', 2)

    recomendaciones = []

    if 'facturas_papel' in analisis and analisis['facturas_papel'].get('total_sospechosas', 0) > 0:
        recomendaciones.append('Revisar las facturas en papel identificadas y requerir a los proveedores el uso de factura electronica cuando sea obligatorio.')

    if 'anotacion' in analisis and analisis['anotacion'].get('facturas_retenidas', 0) > 0:
        recomendaciones.append('Investigar las causas de retencion de facturas en FACe y proceder a su descarga inmediata.')

    if 'validaciones' in analisis and analisis['validaciones'].get('total_incumplimientos', 0) > 0:
        recomendaciones.append('Comunicar a los proveedores los errores detectados en sus facturas para mejorar la calidad de los datos.')

    if 'obligaciones' in analisis and analisis['obligaciones'].get('facturas_3_meses', 0) > 0:
        recomendaciones.append('Establecer un seguimiento especial de las facturas pendientes y coordinar con los organos gestores para acelerar el reconocimiento de obligaciones.')

    recomendaciones.append('Mantener los controles periodicos establecidos y generar informes trimestrales de seguimiento.')

    for recomendacion in recomendaciones:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(recomendacion)

    # Footer
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('\n\n')
    footer.add_run(CONFIGURACION_INFORME['footer']).italic = True
    footer.add_run(f'\n\nGenerado el {datetime.now().strftime("%d/%m/%Y a las %H:%M")}')

    # Guardar en bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


def generar_informe_pdf(datos: dict, analisis: dict) -> bytes:
    """
    Genera un informe ejecutivo en formato PDF con resumen completo
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    story = []
    styles = getSampleStyleSheet()

    # Estilos personalizados
    titulo_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=30,
        alignment=1
    )

    subtitulo_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#333333'),
        spaceAfter=12,
        spaceBefore=20
    )

    seccion_style = ParagraphStyle(
        'SeccionTitle',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=10,
        spaceBefore=15
    )

    alerta_style = ParagraphStyle(
        'Alerta',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#CC0000'),
        spaceBefore=5,
        spaceAfter=5
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=8
    )

    # ============================================
    # PORTADA
    # ============================================
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph(CONFIGURACION_INFORME['titulo'], titulo_style))
    story.append(Paragraph('INFORME EJECUTIVO', subtitulo_style))
    story.append(Spacer(1, 1*cm))

    info_text = f"""
    <b>Entidad:</b> {CONFIGURACION['nombre_entidad']}<br/>
    <b>Ejercicio:</b> {CONFIGURACION['ejercicio_auditado']}<br/>
    <b>Fecha:</b> {datetime.now().strftime('%d/%m/%Y')}
    """
    story.append(Paragraph(info_text, normal_style))
    story.append(PageBreak())

    # ============================================
    # RESUMEN EJECUTIVO
    # ============================================
    story.append(Paragraph('RESUMEN EJECUTIVO', subtitulo_style))

    # Tabla de cifras principales
    total_facturas = len(datos['rcf'])
    facturas_electronicas = len(datos['rcf'][datos['rcf']['es_papel'] == False])
    facturas_papel = len(datos['rcf'][datos['rcf']['es_papel'] == True])

    datos_tabla = [
        ['Concepto', 'Valor'],
        ['Total facturas en RCF', f'{total_facturas:,}'],
        ['Facturas electronicas', f'{facturas_electronicas:,}'],
        ['Facturas en papel', f'{facturas_papel:,}'],
        ['Solicitudes de anulacion', f'{len(datos["anulaciones"]):,}'],
    ]

    tabla = Table(datos_tabla, colWidths=[10*cm, 5*cm])
    tabla.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066CC')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F5F5F5')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC'))
    ]))

    story.append(tabla)
    story.append(Spacer(1, 0.5*cm))

    # ============================================
    # PRINCIPALES HALLAZGOS
    # ============================================
    story.append(Paragraph('PRINCIPALES HALLAZGOS', subtitulo_style))

    hallazgos = []
    alertas = []

    if 'facturas_papel' in analisis:
        papel = analisis['facturas_papel']
        if papel.get('total_sospechosas', 0) > 0:
            alertas.append(f"<b>{papel['total_sospechosas']:,}</b> facturas en papel susceptibles de incumplimiento normativo")

    if 'anotacion' in analisis:
        anot = analisis['anotacion']
        hallazgos.append(f"Tiempo medio de anotacion: <b>{anot.get('tiempo_medio_min', 0):.2f} minutos</b>")
        if anot.get('facturas_retenidas', 0) > 0:
            alertas.append(f"<b>{anot['facturas_retenidas']}</b> facturas retenidas en FACe sin descargar")

    if 'validaciones' in analisis:
        val = analisis['validaciones']
        hallazgos.append(f"Cumplimiento validaciones HAP/1650/2015: <b>{val.get('porcentaje_cumplimiento', 100):.2f}%</b>")
        if val.get('total_incumplimientos', 0) > 0:
            hallazgos.append(f"Incumplimientos detectados: <b>{val['total_incumplimientos']:,}</b>")

    if 'tramitacion' in analisis:
        tram = analisis['tramitacion']
        hallazgos.append(f"Facturas pagadas: <b>{tram.get('facturas_pagadas', 0):,}</b> ({tram.get('importe_pagado', 0):,.0f} EUR)")

    if 'obligaciones' in analisis:
        oblig = analisis['obligaciones']
        if oblig.get('facturas_3_meses', 0) > 0:
            alertas.append(f"<b>{oblig['facturas_3_meses']}</b> facturas pendientes >3 meses ({oblig.get('importe_pendiente', 0):,.0f} EUR)")

    # Mostrar alertas primero
    if alertas:
        story.append(Paragraph('<b>ALERTAS:</b>', alerta_style))
        for alerta in alertas:
            story.append(Paragraph(f"- {alerta}", alerta_style))
        story.append(Spacer(1, 0.3*cm))

    # Mostrar hallazgos
    for hallazgo in hallazgos:
        story.append(Paragraph(f"- {hallazgo}", normal_style))

    story.append(PageBreak())

    # ============================================
    # DETALLE POR SECCION
    # ============================================

    # Facturas en papel
    if 'facturas_papel' in analisis:
        papel = analisis['facturas_papel']
        story.append(Paragraph('FACTURAS EN PAPEL', seccion_style))

        papel_data = [
            ['Metrica', 'Valor'],
            ['Total en papel', f'{papel.get("total_papel", 0):,}'],
            ['Susceptibles incumplimiento', f'{papel.get("total_sospechosas", 0):,}'],
            ['Importe sospechoso', f'{papel.get("importe_sospechoso", 0):,.2f} EUR'],
        ]

        tabla = Table(papel_data, colWidths=[10*cm, 5*cm])
        tabla.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FF6B35')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC')),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FFF5F0')),
        ]))
        story.append(tabla)
        story.append(Spacer(1, 0.3*cm))

    # Tiempos de anotacion
    if 'anotacion' in analisis:
        anot = analisis['anotacion']
        story.append(Paragraph('TIEMPOS DE ANOTACION', seccion_style))

        anot_data = [
            ['Metrica', 'Valor'],
            ['Facturas retenidas', f'{anot.get("facturas_retenidas", 0):,}'],
            ['Tiempo medio', f'{anot.get("tiempo_medio_min", 0):.2f} min'],
            ['Tiempo maximo', f'{anot.get("tiempo_max_min", 0):.2f} min'],
        ]

        tabla = Table(anot_data, colWidths=[10*cm, 5*cm])
        tabla.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#28A745')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC')),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F0FFF0')),
        ]))
        story.append(tabla)
        story.append(Spacer(1, 0.3*cm))

    # Validaciones
    if 'validaciones' in analisis:
        val = analisis['validaciones']
        story.append(Paragraph('VALIDACIONES HAP/1650/2015', seccion_style))

        val_data = [
            ['Metrica', 'Valor'],
            ['Facturas validadas', f'{val.get("total_facturas_validadas", 0):,}'],
            ['Incumplimientos', f'{val.get("total_incumplimientos", 0):,}'],
            ['Cumplimiento', f'{val.get("porcentaje_cumplimiento", 100):.2f}%'],
        ]

        tabla = Table(val_data, colWidths=[10*cm, 5*cm])
        tabla.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6F42C1')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC')),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F0FF')),
        ]))
        story.append(tabla)
        story.append(Spacer(1, 0.3*cm))

    # Obligaciones
    if 'obligaciones' in analisis:
        oblig = analisis['obligaciones']
        story.append(Paragraph('OBLIGACIONES PENDIENTES', seccion_style))

        oblig_data = [
            ['Metrica', 'Valor'],
            ['Facturas >3 meses', f'{oblig.get("facturas_3_meses", 0):,}'],
            ['Importe pendiente', f'{oblig.get("importe_pendiente", 0):,.2f} EUR'],
            ['Dias medio pendiente', f'{oblig.get("dias_medio_pendiente", 0):.0f} dias'],
        ]

        tabla = Table(oblig_data, colWidths=[10*cm, 5*cm])
        tabla.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#DC3545')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC')),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FFF0F0')),
        ]))
        story.append(tabla)

    # Footer
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph('<hr/>', normal_style))
    footer_text = f"<i>{CONFIGURACION_INFORME['footer']}</i><br/><br/>Generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}"
    story.append(Paragraph(footer_text, normal_style))

    # Construir PDF
    doc.build(story)
    buffer.seek(0)

    return buffer.getvalue()
